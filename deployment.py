# -*- coding: utf-8 -*-
"""
Created on Thu Mar  9 08:46:48 2023

@author: ashwi
"""
# pip install python-docx
# python -m spacy download en
# !pip install spacy==(last version of spacy(3.2.4)) 

import streamlit as st 
import streamlit.components.v1 as stc
import spacy
import pickle 
import random
from spacy import displacy
import docx
from spacy.lang.en.stop_words import  STOP_WORDS
import string
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer,CountVectorizer
from sklearn.model_selection import train_test_split
from sklearn.metrics import accuracy_score, classification_report, confusion_matrix
import PyPDF2
import re
import nltk
from spacy.matcher import Matcher
#nltk.download('stopwords')
#nltk.download('wordnet')
import pickle
from pickle import dump
from pickle import load
from nltk.stem.porter import PorterStemmer
from nltk.stem import WordNetLemmatizer
from sklearn.metrics import confusion_matrix,roc_auc_score,accuracy_score,classification_report
from sklearn.multiclass import OneVsRestClassifier
from sklearn import metrics
from sklearn.metrics import accuracy_score
from pandas.plotting import scatter_matrix
from sklearn.ensemble import RandomForestClassifier
from sklearn.pipeline import Pipeline
import pdfplumber
import spacy_transformers
import en_core_web_trf
nlp = spacy.load("en_core_web_trf")

from sklearn.preprocessing import LabelEncoder
le_encoder=LabelEncoder()


        

st.title('Model Deployment ')
st.title('Resumes Classification')

st.header('User Input Resume')
st.sidebar.subheader('File Description')
st.sidebar.subheader('About')

html_temp = """
    <div style="background-color:white;padding:7px">
    <h2 style="color:black;text-align:center;"> Document Classifier</h2>
    </div>
    """
    



def readtxt(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


def readpdf(filename):
    pdf=pdfplumber.open(filename)
    pages = pdf.pages[0]
    text=pages.extract_text()
    return text

# Defining the SkillSet
skill_set=[  't-sql', 'sas', 'r', 'python', 'mariadb',
            'msexcel', 'tableau', 'xml', 'xslt', 'eib',
           'oracle', 'peoplesoft', 'sql', 'hcm', 'fcm',
           'msbi', 'html', 'css3', 'css',
           'xml', 'javascript', 'json', 'reactjs', 'nodejs',"java","c","c++",
              'workday', 'hcm', 'eib', 'picof','workday hcm',
                        'workday studio','nnbound/outbound integrations',
                        'peoplesoft', 'pia','ccb','birt','peci','ccw','pum','people tools',
                        'peoplesoft implementation','peoplesoft components',
                        'peoplesoft dba','peoplesoft admin','peoplesoft admin/dba','peopleSoft fscm', 
                        'peopletoolsupgrade','peopletools upgrade','process scheduler servers',
                        'peoplesoft hrms','peopleSoft consultant','peopledoft cloud',
                        'PeopleSoft migrations','eoplesoft Testing Framework','pure internet architecture',
                        'sql','sql server', 'ms sql server','msbi', 'sql developer', 'ssis','ssrs',
                        'ssms','t-sql','tsql','Razorsql', 'razor sql','triggers','powerbi','power bi',
                        'oracle sql', 'pl/sql', 'pl\sql','oracle', 'oracle 11g','oledb','cte','ddl',
                        'dml','etl','mariadb','maria db','reactjs', 'react js','React JS','ReactJS', 'react js developer', 'html', 
                        'css3','xml','javascript','html5','boostrap','jquery', 'redux','php', 'node js',
                        'nodejs','apache','netbeans','nestjs','nest js','react developer','react hooks',
                        'jenkins','datascience','python','data science','ML','machine learning','Machine Learning',
                        'Machine learning','AI','ai','Artificial Inteligence','artifitial ineligence','Modelling',
                        'Big Data','Bigdata','BigData','bigdata','AWS','Cloud Environments','oracle','Oracle']

#Skill Set Extraction Function

def extract_skills(resume_text):
    nlp_text = nlp(resume_text)
    noun_chunks = nlp_text.noun_chunks
    # removing stop words and implementing word tokenization
    tokens = [token.text for token in nlp_text if not token.is_stop]
    
    
     
    
    
    skills = skill_set
    
    skillset = []
    
    # check for one-grams
    for token in tokens:
        if token in skills:
            skillset.append(token)
    
    # check for bi-grams and tri-grams
    for token in noun_chunks:
        token = token.text.lower().strip()
        if token in skills:
            skillset.append(token)
    
    return [i.capitalize() for i in set([i.lower() for i in skillset])]

def uniquify(string):
    output = []
    seen = set()
    for word in string.split():
        if word not in seen:
            output.append(word)
            seen.add(word)
    return ' '.join(output)

# Defining the Category name

category_name=[  'peoplesoft','Peoplesoft','ReactDeveloper','Reactjs developer','workday resumes','workday',
                 'workday studio','internship','Internship','sql','sql developer','SQL Developer','ui developer','workday consultant']

#Category name Extraction Function

def extract_category(resume_text):
    nlp_text = nlp(resume_text)
    noun_chunks = nlp_text.noun_chunks
    # removing stop words and implementing word tokenization
    tokens = [token.text for token in nlp_text if not token.is_stop]


    category = category_name
    
    categoryname = [ ]

    # check for one-grams
    for token in tokens:
        if token in category:
            categoryname.append(token)

# check for bi-grams and tri-grams
    for token in noun_chunks:
        token = token.text.lower().strip()
        if token in category:
            categoryname.append(token)
    
    return [i.capitalize() for i in set([i.lower() for i in categoryname])]


def main():
    st.markdown(html_temp,unsafe_allow_html=True)
    docx_file = st.file_uploader("Upload Document", type=["pdf","docx","txt"])
    if st.button("Process"):
        if docx_file is not None:
            file_details = {"filename":docx_file.name, "filetype":docx_file.type,
                                        "filesize":docx_file.size}
            st.sidebar.write(file_details)
            if docx_file.type == "text/plain":
                
            	# Read as string (decode bytes to string)
            	raw_text = str(docx_file.read(),"utf-8")
                

            elif docx_file.type == "application/pdf":
                raw_text =readpdf(docx_file)

            else:
                raw_text = readtxt(docx_file) 
                
        return raw_text
                
            	

                                                    
def file_input():
    fi11=main()
    li=[]
    li.append(fi11)
    data = {'resume':li}
    features = pd.DataFrame(data,index = [0])
    
    return features
    
        

				
    
    
     

df = file_input()


def text_preprocessing(df):
    
    html_temp = """
        <div style="background-color:white;padding:5px">
        <h2 style="color:black;text-align:center;">Extracted Skills</h2>
        </div>
        """
    st.markdown(html_temp,unsafe_allow_html=True)
    
    try:
        clean = []
        lz = WordNetLemmatizer()
        for i in range(df.shape[0]):
            review = re.sub(
                '(@[A-Za-z0-9]+)|([^0-9A-Za-z \t])|(\w+:\/\/\S+)|^rt|http.+?"',
                " ",
                df["resume"].iloc[i],
            )
            review = re.sub(r"[0-9]+", " ", review) # Remove Numbers
            review = review.lower()
            review = review.split()
            lm = WordNetLemmatizer()
            review = [ lz.lemmatize(word) for word in review if word not in STOP_WORDS]
            review = " ".join(review)
            clean.append(review)
  
    except:
        pass
  
   


    df["Clean_Resume"] = clean
    df["Clean_Resume"]=df["Clean_Resume"].apply(uniquify)

    st.write(df["Clean_Resume"])
    df["Skills"]=df["Clean_Resume"].apply(extract_skills)
    df["Category"]=df["Clean_Resume"].apply(extract_category)

    st.subheader('Skills Extracted From Resume')
    st.write(df["Skills"])

    st.subheader('Category Extracted From Resume')
    st.write(df["Category"])


    # Creating Resume Data Frame
    resume_data=pd.DataFrame()
    resume_data["Resume"]=df["Clean_Resume"]


    from PIL import Image
    image = Image.open('C:\\Users\\alank\\project2deployment\\Screenshot 2022-09-10 19.13.11.png')

    st.image(image, caption='pie chart of dataset')

    image = Image.open('C:\\Users\\alank\\project2deployment\\Screenshot 2022-09-10 19.13.35.png')

    st.image(image, caption='wordcloud image')

  # Function to extract text from resume
def getText(filename):
      
    # Create empty string 
    fullText = ''
    if filename.endswith('.docx'):
        doc = docx.Document(filename)
        
        for para in doc.paragraphs:
            fullText = fullText + para.text
            
           
    elif filename.endswith('.pdf'):  
        with open(filename, "rb") as pdf_file:
            pdoc = PyPDF2.PdfFileReader(filename)
            number_of_pages = pdoc.getNumPages()
            page = pdoc.pages[0]
            page_content = page.extractText()
             
        for paragraph in page_content:
            fullText =  fullText + paragraph
            
    else:
        import aspose.words as aw
        output = aw.Document()
        # Remove all content from the destination document before appending.
        output.remove_all_children()
        input = aw.Document(filename)
        # Append the source document to the end of the destination document.
        output.append_document(input, aw.ImportFormatMode.KEEP_SOURCE_FORMATTING)
        output.save("Output.docx");
        doc = docx.Document('Output.docx')
        
        for para in doc.paragraphs:
            fullText = fullText + para.text
        fullText = fullText[79:]
         
    return (fullText)




    #  Vectorisation


    requiredText = resume_data["Resume"].values
    #word_vectorizer = CountVectorizer(analyzer='word',token_pattern=r'\w{1,}',stop_words = 'english')
    #word_vectorizer.fit(requiredText)
    #WordFeatures = word_vectorizer.fit_transform(requiredText)
    
    return   requiredText               #WordFeatures

try:
  df11=text_preprocessing(df)



    
    
    
  resume_data=pd.read_csv("C:/Users/ashwi/Desktop/NLP project/Resume.csv")
  resume_data["Encoded_Skill"]=le_encoder.fit_transform(resume_data["Category"])


   
 
  requiredText1 = resume_data["Resume"]
  requiredTarget1 = resume_data["Encoded_Skill"].values
#vectorizer = CountVectorizer(analyzer='word',token_pattern=r'\w{1,}',stop_words = 'english')
#vectorizer.fit(requiredText1)
#WordFeatures1 = vectorizer.fit_transform(requiredText1)
    
  X_train,X_test,y_train,y_test = train_test_split(requiredText1,requiredTarget1,
                                                 stratify=requiredTarget1,random_state=22, test_size=0.20)



  pipe_lr = Pipeline(steps=[('cv',CountVectorizer(analyzer='word',token_pattern=r'\w{1,}',stop_words = 'english')),
                          ('rf',RandomForestClassifier(n_estimators=250,criterion='entropy',max_depth=2,
                                                      max_features= 0.8 ,
                                                      random_state=None, class_weight="balanced"))])

#lr = RandomForestClassifier(n_estimators=100,criterion='entropy',max_depth=5,
                            #max_features= None ,random_state=None, class_weight="balanced")
  pipe_lr.fit(X_train, y_train)
    

  lr_prediction = pipe_lr.predict(df11)
  lr_score = pipe_lr.score(X_test,y_test)
  y_pred=lr_prediction

  proba=pipe_lr.predict_proba(X_test)

  html_temp = """
    <div style="background-color:green;padding:5px">
    <h2 style="color:white;text-align:center;">Different Categories Of the Target</h2>
    </div>
    """
  st.markdown(html_temp,unsafe_allow_html=True)

#st.subheader('Probabilities')
#st.write(proba)


  cls=pipe_lr.classes_
  
  col1, col2 = st.columns(2)

  col1.subheader('Classes')
  col1.write(resume_data.category.unique())




  result= le_encoder.inverse_transform(lr_prediction)
  


  import numpy as np
  import matplotlib.pyplot as plt

  st.set_option('deprecation.showPyplotGlobalUse', False)

  def plot():
      
      import numpy as np
      from matplotlib.gridspec import GridSpec
      targetCounts = resume_data.category.value_counts()
      targetLabels  = resume_data.category.unique()
      # Make square figures and axes
      plt.figure(1, figsize=(25,25))
      the_grid = GridSpec(2, 2)


      cmap = plt.get_cmap('plasma')
      colors = [cmap(i) for i in np.linspace(0, 1, 6)]
      plt.subplot(the_grid[0, 1], aspect=1, title='CATEGORY DISTRIBUTION')
      plt.style.use('seaborn-dark-palette')

      source_pie = plt.pie(targetCounts, labels=targetLabels, autopct='%1.1f%%', shadow=True, colors=colors)
      plt.show()
    
    
  p=plot()

  html_temp = """
    <div style="background-color:blue;padding:5px">
    <h2 style="color:white;text-align:center;">Prediction Result</h2>
    </div>
    """

  st.markdown(html_temp,unsafe_allow_html=True)


  st.subheader('Prediction Result')
#st.write(y_pred)
  st.write(result)

  col2.subheader('Category Plot')
  col2.pyplot(p)
  
  
  
  
  st.sidebar.text("'Project Done By Group51\n @ Mr N Ashwin siddhartha & Others  '")
  
  

    

    

  

except:
    
    pass
